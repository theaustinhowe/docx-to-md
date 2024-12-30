from docx import Document
import re

# Function to convert text to Markdown-safe format
def escape_markdown(text):
    # Escape Markdown special characters
    return re.sub(r'([#*_`~|])', r'\\\1', text)

# Function to parse document paragraphs to Markdown
def parse_docx_to_markdown(docx_file, output_file):
    # Load the .docx file
    doc = Document(docx_file)

    with open(output_file, 'w', encoding='utf-8') as md_file:
        for paragraph in doc.paragraphs:
            # Check for heading styles and apply appropriate Markdown syntax
            if paragraph.style.name.startswith("Heading 1"):
                md_file.write(f"# {escape_markdown(paragraph.text)}\n\n")
            elif paragraph.style.name.startswith("Heading 2"):
                md_file.write(f"## {escape_markdown(paragraph.text)}\n\n")
            elif paragraph.style.name.startswith("Heading 3"):
                md_file.write(f"### {escape_markdown(paragraph.text)}\n\n")
            elif paragraph.style.name.startswith("Heading 4"):
                md_file.write(f"#### {escape_markdown(paragraph.text)}\n\n")
            elif paragraph.style.name.startswith("Heading 5"):
                md_file.write(f"##### {escape_markdown(paragraph.text)}\n\n")
            elif paragraph.style.name.startswith("Heading 6"):
                md_file.write(f"###### {escape_markdown(paragraph.text)}\n\n")
            else:
                # Check for bold and italic text
                if paragraph.runs:
                    line = ""
                    for run in paragraph.runs:
                        text = escape_markdown(run.text)
                        if run.bold and run.italic:
                            line += f"***{text}***"
                        elif run.bold:
                            line += f"**{text}**"
                        elif run.italic:
                            line += f"*{text}*"
                        else:
                            line += text
                    md_file.write(f"{line}\n\n")
                else:
                    if paragraph.text.strip():
                        md_file.write(f"{escape_markdown(paragraph.text)}\n\n")

        # Process tables
        for table in doc.tables:
            # Write table headers
            headers = [escape_markdown(cell.text.strip()) for cell in table.rows[0].cells]
            md_file.write(f"| {' | '.join(headers)} |\n")
            md_file.write(f"| {' | '.join(['---' for _ in headers])} |\n")

            # Write table rows
            for row in table.rows[1:]:
                row_data = [escape_markdown(cell.text.strip()) for cell in row.cells]
                md_file.write(f"| {' | '.join(row_data)} |\n")
            md_file.write(f"\n")

        # Process images (if any)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                md_file.write(f"![Image]({rel.target_ref})\n\n")

# Example usage
input_docx = "Work History.docx"  # Replace with your .docx file path
output_md = "output.md"         # Replace with your desired output file path

parse_docx_to_markdown(input_docx, output_md)

print(f"Markdown file has been saved to {output_md}")
